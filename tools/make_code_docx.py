# -*- coding: utf-8 -*-
"""
Generate a Word (.docx) document containing project code snippets rendered
in a dark IDE-style code panel: syntax coloring + a line-number gutter.
Font: Arial (as requested).

Layout note: each snippet is a SINGLE 2-cell table (one cell of line numbers,
one cell of code), with every line as a line-broken paragraph inside the cell.
This imports cleanly into Google Docs, unlike a one-row-per-line table which
explodes into gappy cells when pasted/imported.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pygments.lexers import get_lexer_by_name
from pygments.token import Token

# ---------------------------------------------------------------------------
# VS Code "Light+" inspired palette (readable on a white background)
# ---------------------------------------------------------------------------
PANEL_BG   = "FFFFFF"   # editor background (white)
GUTTER_BG  = "F3F3F3"   # line-number gutter (very light gray)
LINE_NO    = "6E7781"   # line numbers (gray)
DEFAULT_FG = "000000"   # default code text

COLOR_MAP = [
    (Token.Comment,            "008000"),
    (Token.Keyword,            "0000FF"),
    (Token.Operator.Word,      "0000FF"),
    (Token.Name.Builtin,       "0000FF"),
    (Token.Name.Builtin.Pseudo,"0000FF"),
    (Token.Name.Class,         "267F99"),
    (Token.Name.Function,      "795E26"),
    (Token.Name.Decorator,     "795E26"),
    (Token.Name.Tag,           "800000"),
    (Token.Name.Attribute,     "E50000"),
    (Token.Name.Variable,      "001080"),
    (Token.Name.Constant,      "0070C1"),
    (Token.Literal.String,     "A31515"),
    (Token.Literal.Number,     "098658"),
    (Token.Operator,           "000000"),
    (Token.Punctuation,        "000000"),
    (Token.Name,               "001080"),
]

def token_color(ttype):
    t = ttype
    while t is not None:
        for base, color in COLOR_MAP:
            if t is base:
                return color
        t = t.parent
    return DEFAULT_FG

# ---------------------------------------------------------------------------
# low-level XML helpers
# ---------------------------------------------------------------------------
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement('w:tcMar')
    for name, val in (('top', top), ('bottom', bottom), ('start', left),
                      ('end', right), ('left', left), ('right', right)):
        node = OxmlElement('w:' + name)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        margins.append(node)
    tcPr.append(margins)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'none')
        e.set(qn('w:sz'), '0')
        e.set(qn('w:space'), '0')
        borders.append(e)
    tblPr.append(borders)

def style_run(run, color_hex, bold=False):
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    rFonts.set(qn('w:cs'), 'Arial')

def _line_para(cell, first):
    """Tight, zero-spacing paragraph in the cell (reuse the 1st, add the rest)."""
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    return p

# ---------------------------------------------------------------------------
# build a code panel (one table) for a snippet
# ---------------------------------------------------------------------------
def tokenize_lines(code, lang):
    """Return list of lines, each a list of (text, tokentype)."""
    opts = dict(stripnl=False, stripall=False)
    if lang == 'php':
        # snippets are bare PHP (no <?php tag) -> highlight inline code
        opts['startinline'] = True
    lexer = get_lexer_by_name(lang, **opts)
    lines = [[]]
    for ttype, value in lexer.get_tokens(code):
        parts = value.split('\n')
        for i, part in enumerate(parts):
            if i > 0:
                lines.append([])
            if part:
                lines[-1].append((part, ttype))
    if lines and not lines[-1]:
        lines.pop()
    return lines

# Left geometry matched to v3 (the table layout the user preferred):
#   v3 line-number right edge ~= 480 twips ; v3 code start ~= 780 twips
NUM_TAB  = Twips(480)   # right-aligned tab: line number ends here
CODE_TAB = Twips(780)   # left tab + hanging indent: code starts here

def add_code_panel(doc, code, lang, start_line):
    """Render each code line as its own paragraph (NOT a table). A right-aligned
    tab stop holds the line number in a gutter; a left tab stop starts the code.
    Paragraphs flow and break across pages naturally, so there is no 'dead space'
    in Google Docs (a single big table row cannot split across a page there).
    Left-side spacing matches version 3."""
    lines = tokenize_lines(code, lang)

    for idx, line_tokens in enumerate(lines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        # hanging indent so wrapped long lines align under the code, not the gutter
        pf.left_indent = CODE_TAB
        pf.first_line_indent = -CODE_TAB
        pf.tab_stops.add_tab_stop(NUM_TAB, WD_TAB_ALIGNMENT.RIGHT)   # number
        pf.tab_stops.add_tab_stop(CODE_TAB, WD_TAB_ALIGNMENT.LEFT)   # code

        # line number, right-aligned in the gutter
        r = p.add_run()
        r.add_tab()
        style_run(r, LINE_NO)
        style_run(p.add_run(str(start_line + idx)), LINE_NO)

        # tab over to the code column
        sep = p.add_run()
        sep.add_tab()
        style_run(sep, DEFAULT_FG)

        # code
        if not line_tokens:
            style_run(p.add_run(' '), DEFAULT_FG)  # keep blank lines alive
        else:
            for text, ttype in line_tokens:
                style_run(p.add_run(text), token_color(ttype))

# ---------------------------------------------------------------------------
# snippet definitions
# ---------------------------------------------------------------------------
SNIPPETS = [
    {
        "title": "app/Services/DocumentQueryService.php",
        "lang": "php",
        "start": 234,
        "code": r'''        $embeddings = DocumentEmbedding::with('document:doc_id,title,status')
            ->whereHas('document', function($query) {
                $query->where('status', 'active');
            })
            ->orderBy('doc_id')
            ->orderBy('chunk_index')
            ->get()
            ->map(function ($embedding) {
                return [
                    'embedding_id' => $embedding->embedding_id,
                    'doc_id' => $embedding->doc_id,
                    'document_title' => $embedding->document ? $embedding->document->title : 'Unknown',
                    'chunk_index' => $embedding->chunk_index,
                    'chunk_text' => $embedding->chunk_text,
                    'embedding_vector' => json_decode($embedding->embedding_vector),
                    'created_at' => $embedding->created_at
                ];
            });

        return [
            'success' => true,
            'total_embeddings' => $embeddings->count(),
            'data' => $embeddings
        ];
    }

    /**
     * Get latest uploaded document by user
     */
    public function getLatestProcessingDocument(int $userId): ?Document
    {
        return Document::where('created_by', $userId)
            ->whereIn('status', ['processing', 'processed'])
            ->latest('created_at')
            ->first();
    }

    /**
     * AI semantic search using Groq API
     * Returns doc IDs that semantically match the search term.
     * Falls back to empty array (SQL-only results) if API fails or times out.
     */
    private function getSemanticSearchResults(string $searchTerm): array
    {
        try {
            $searchApiKey = env('GROQ_SEARCH_API_KEY');
            if (empty($searchApiKey)) {
                return [];
            }''',
    },
    {
        "title": "scanner_service/server.js",
        "lang": "javascript",
        "start": 91,
        "code": r'''    const command = `"${NAPS2_PATH}" --driver ${driver} --device "${safeDevice}" --source Feeder -o "${outputPath}" -f -v`;

    console.log(`Executing: ${command}`);

    // Execute NAPS2 Console
    exec(command, async (error, stdout, stderr) => {
        if (error) {
            console.error('NAPS2 Execution Error:', error.message);

            const errorDetails = stderr || error.message;
            let userMessage = 'Scanner execution failed.';

            // Common NAPS2 error patterns
            if (errorDetails.includes('No profiles')) {
                userMessage = 'No scan profiles found. Please open NAPS2 and create a profile.';
            } else if (errorDetails.includes('No device')) {
                userMessage = 'No scanner device detected or profiles are invalid.';
            }

            return res.status(500).json({
                success: false,
                message: userMessage,
                details: errorDetails,
                stderr: stderr
            });
        }

        console.log('NAPS2 Output:', stdout);
        console.log('Scan generated at:', outputPath);

        // Verify file exists
        if (!fs.existsSync(outputPath)) {
            return res.status(500).json({
                success: false,
                message: 'Scan command finished but NO file was generated.',
                details: 'Check if the scanner actually scanned a page.'
            });
        }

        // Upload to Main Application (Laravel)
        try {
            console.log('Uploading to backend...');
            const formData = new FormData();
            formData.append('file', fs.createReadStream(outputPath));

            // Reads LARAVEL_URL from .env file - change this when hosted
            const API_URL = `${process.env.LARAVEL_URL || 'http://127.0.0.1:8000'}/api/scanner/upload`;

            const uploadResponse = await axios.post(API_URL, formData, {
                headers: { ...formData.getHeaders() },
                maxContentLength: Infinity,
                maxBodyLength: Infinity
            });''',
    },
    {
        "title": "aiservice/ai_bridge/ai_service.py",
        "lang": "python",
        "start": 190,
        "code": r'''            with ThreadPoolExecutor(max_workers=3) as executor:
                future_title = executor.submit(self._extract_title, text)
                future_desc = executor.submit(self._generate_description, text)
                future_remarks = executor.submit(self._generate_remarks, text)

                try:
                    title = future_title.result(timeout=60)
                except Exception as e:
                    logger.error(f"Parallel title generation failed: {str(e)}")

                try:
                    description = future_desc.result(timeout=60)
                except Exception as e:
                    logger.error(f"Parallel description generation failed: {str(e)}")

                try:
                    remarks = future_remarks.result(timeout=60)
                except Exception as e:
                    logger.error(f"Parallel remarks generation failed: {str(e)}")

            logger.info(f"All 3 AI tasks completed in parallel")

            return {
                'suggested_title': title,
                'suggested_description': description,
                'ai_remarks': remarks
            }''',
    },
    {
        "title": "app/Http/Middleware/RedirectBasedOnRole.php",
        "lang": "php",
        "start": 17,
        "code": r'''    public function handle(Request $request, Closure $next, string $role): Response
    {
        if (!Auth::check()) {
            return redirect('/login');
        }

        $user = Auth::user();

        // Check if user has the required role
        if ($user->role !== $role) {
            // Redirect to appropriate dashboard based on actual role
            if ($user->role === 'admin') {
                return redirect('/admin/dashboard');
            } elseif ($user->role === 'staff') {
                return redirect('/staff/dashboard');
            }
        }

        return $next($request);
    }''',
    },
    {
        "title": "app/Http/Controllers/AIAssistantController.php",
        "lang": "php",
        "start": 333,
        "code": r'''            try {
                if ($primaryService === 'groq') {
                    Log::info('Attempting Groq API for chat...');
                    $aiResponse = $this->callGroqAPI($request->message, $conversationId, $documentContext);
                } else {
                    Log::info('Attempting local AI service...');
                    $aiResponse = $this->callLocalAIService($request->message, $conversationId, $documentContext);
                }
            } catch (\Exception $e) {
                Log::warning("Primary AI service ({$primaryService}) failed: " . $e->getMessage());

                // Automatic fallback to alternative service
                try {
                    if ($primaryService === 'groq') {
                        Log::info('Groq failed, falling back to local AI service...');
                        $aiResponse = $this->callLocalAIService($request->message, $conversationId, $documentContext);
                    } else {
                        Log::info('Local AI failed, falling back to Groq API...');
                        $aiResponse = $this->callGroqAPI($request->message, $conversationId, $documentContext);
                    }
                } catch (\Exception $fallbackError) {
                    Log::error('Both AI services failed', [
                        'primary_error' => $e->getMessage(),
                        'fallback_error' => $fallbackError->getMessage()
                    ]);
                    throw new \Exception('All AI services are currently unavailable. Please ensure either the local AI server is running or you have internet connection for Groq API.');
                }
            }''',
    },
]

# ---------------------------------------------------------------------------
# assemble document
# ---------------------------------------------------------------------------
def main():
    doc = Document()

    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(11)

    title = doc.add_heading('Legal Arch AIU - Code Snippets', level=0)
    for run in title.runs:
        run.font.name = 'Arial'

    for snip in SNIPPETS:
        h = doc.add_heading(snip["title"], level=2)
        for run in h.runs:
            run.font.name = 'Arial'
        add_code_panel(doc, snip["code"], snip["lang"], snip["start"])
        doc.add_paragraph()  # spacer

    out = r'd:\capstone\Legal_Arch_aiu\Legal_Arch_AIU_Code_Snippets_v5.docx'
    doc.save(out)
    print('Saved:', out)

if __name__ == '__main__':
    main()
