"""Convert USER_MANUAL.md to a formatted Word document."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = "USER_MANUAL.md"
OUT = "USER_MANUAL.docx"

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def add_runs_with_bold(paragraph, text):
    """Add text to a paragraph, rendering **bold** and *italic* segments."""
    # Split on bold first
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


with open(SRC, "r", encoding="utf-8") as f:
    lines = f.read().split("\n")

i = 0
n = len(lines)
while i < n:
    line = lines[i]
    stripped = line.strip()

    # Blank line
    if not stripped:
        i += 1
        continue

    # Horizontal rule
    if stripped == "---":
        i += 1
        continue

    # Table detection: a line of | ... | followed by a |---| separator
    if stripped.startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:\-|]+\|\s*$", lines[i + 1]):
        header = parse_table_row(lines[i])
        i += 2  # skip header + separator
        rows = []
        while i < n and lines[i].strip().startswith("|"):
            rows.append(parse_table_row(lines[i]))
            i += 1
        table = doc.add_table(rows=1, cols=len(header))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for j, h in enumerate(header):
            hdr_cells[j].text = ""
            p = hdr_cells[j].paragraphs[0]
            add_runs_with_bold(p, h)
            for run in p.runs:
                run.bold = True
        for row in rows:
            cells = table.add_row().cells
            for j, val in enumerate(row):
                if j < len(cells):
                    cells[j].text = ""
                    add_runs_with_bold(cells[j].paragraphs[0], val)
        doc.add_paragraph()
        continue

    # Headings
    if stripped.startswith("#### "):
        p = doc.add_heading(level=4)
        add_runs_with_bold(p, stripped[5:])
        i += 1
        continue
    if stripped.startswith("### "):
        p = doc.add_heading(level=3)
        add_runs_with_bold(p, stripped[4:])
        i += 1
        continue
    if stripped.startswith("## "):
        p = doc.add_heading(level=2)
        add_runs_with_bold(p, stripped[3:])
        i += 1
        continue
    if stripped.startswith("# "):
        p = doc.add_heading(level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs_with_bold(p, stripped[2:])
        i += 1
        continue

    # Blockquote / Note
    if stripped.startswith(">"):
        text = stripped.lstrip(">").strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        add_runs_with_bold(p, text)
        for run in p.runs:
            run.italic = True
        i += 1
        continue

    # Numbered / lettered list (1.  a.  etc.)
    m = re.match(r"^(\s*)([0-9]+|[a-z]|[ivx]+)\.\s+(.*)$", line)
    if m:
        indent = len(m.group(1))
        p = doc.add_paragraph(style="List Number" if indent == 0 else "List Number 2")
        add_runs_with_bold(p, m.group(3))
        i += 1
        continue

    # Bullet list
    m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
    if m:
        indent = len(m.group(1))
        style = "List Bullet" if indent < 2 else "List Bullet 2"
        p = doc.add_paragraph(style=style)
        add_runs_with_bold(p, m.group(2))
        i += 1
        continue

    # Plain paragraph
    p = doc.add_paragraph()
    add_runs_with_bold(p, stripped)
    i += 1

BLACK = RGBColor(0, 0, 0)


def force_black(paragraph):
    for run in paragraph.runs:
        run.font.color.rgb = BLACK


# Force all body paragraphs (incl. headings) to black
for p in doc.paragraphs:
    force_black(p)

# Force all table cells to black
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                force_black(p)

doc.save(OUT)
print(f"Saved {OUT}")
