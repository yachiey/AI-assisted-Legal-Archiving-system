"""Generate the Data Dictionary Word document for the AI-Powered Legal
Document Management System. Reflects the actual database schema in
database/migrations/ and the Eloquent models in app/Models/.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HEADER_FILL = "1F3864"  # dark blue
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
ZEBRA_FILL = "F2F2F2"


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "808080")
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    return p


def add_paragraph(doc, text, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    table.allow_autofit = False

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = HEADER_TEXT
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(hdr_cells[i], HEADER_FILL)
        set_cell_borders(hdr_cells[i])

    # Body rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_borders(cells[c_idx])
            if r_idx % 2 == 1:
                shade_cell(cells[c_idx], ZEBRA_FILL)

    # Column widths (Cm) — tuned for an A4/Letter page
    widths = [Cm(3.6), Cm(2.7), Cm(3.8), Cm(1.2), Cm(1.8), Cm(4.5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i] if i < len(widths) else Cm(3.0)

    return table


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)


HEADERS = ["Attribute", "Data Type", "Constraints", "Null", "Default", "Description"]


TABLES = [
    {
        "caption": "Table 4.1. Data Dictionary — users",
        "intro": "Stores all system user accounts and their role-based permission flags.",
        "rows": [
            ["user_id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "No", "—", "Unique identifier of the user."],
            ["lastname", "VARCHAR(255)", "—", "No", "—", "User's surname."],
            ["firstname", "VARCHAR(255)", "—", "No", "—", "User's given name."],
            ["middle_name", "VARCHAR(255)", "—", "Yes", "NULL", "User's middle name."],
            ["email", "VARCHAR(255)", "UNIQUE", "No", "—", "Login email address."],
            ["password", "VARCHAR(255)", "—", "No", "—", "Bcrypt-hashed password."],
            ["profile_picture", "VARCHAR(255)", "—", "Yes", "NULL", "Path to the user's avatar image."],
            ["role", "ENUM('admin','staff')", "—", "No", "'staff'", "Role-based access classification."],
            ["status", "ENUM('active','inactive')", "—", "No", "'active'", "Account activation status."],
            ["can_view", "BOOLEAN", "—", "No", "true", "Permission flag for viewing documents."],
            ["can_upload", "BOOLEAN", "—", "No", "false", "Permission flag for uploading documents."],
            ["can_edit", "BOOLEAN", "—", "No", "false", "Permission flag for editing documents."],
            ["can_delete", "BOOLEAN", "—", "No", "false", "Permission flag for deleting documents."],
            ["can_archive", "BOOLEAN", "—", "No", "false", "Permission flag for archiving documents."],
            ["remember_token", "VARCHAR(100)", "—", "Yes", "NULL", "Laravel \"remember me\" session token."],
            ["created_at", "TIMESTAMP", "—", "Yes", "NULL", "Record creation timestamp."],
            ["updated_at", "TIMESTAMP", "—", "Yes", "NULL", "Record last-update timestamp."],
        ],
    },
    {
        "caption": "Table 4.2. Data Dictionary — folders",
        "intro": "Stores the hierarchical folder structure used to organize the document archive.",
        "rows": [
            ["folder_id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "No", "—", "Unique identifier of the folder."],
            ["folder_name", "VARCHAR(255)", "INDEX", "No", "—", "Display name of the folder."],
            ["folder_path", "VARCHAR(255)", "—", "No", "—", "Full hierarchical path of the folder."],
            ["parent_folder_id", "BIGINT UNSIGNED", "FK → folders(folder_id), ON DELETE CASCADE, INDEX", "Yes", "NULL", "Self-reference to the parent folder; NULL for root folders."],
            ["created_by", "BIGINT UNSIGNED", "FK → users(user_id), ON DELETE CASCADE", "No", "—", "User who created the folder."],
            ["folder_type", "VARCHAR(255)", "—", "No", "—", "Classification of the folder (e.g., legal category)."],
            ["created_at", "TIMESTAMP", "—", "Yes", "NULL", "Record creation timestamp."],
            ["updated_at", "TIMESTAMP", "INDEX", "Yes", "NULL", "Record last-update timestamp."],
        ],
    },
    {
        "caption": "Table 4.3. Data Dictionary — documents",
        "intro": "Stores metadata for every uploaded legal document. Composite index (folder_id, status) supports filtered folder queries.",
        "rows": [
            ["doc_id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "No", "—", "Unique identifier of the document."],
            ["document_ref_id", "VARCHAR(255)", "UNIQUE", "Yes", "NULL", "Human-readable reference code for the document."],
            ["title", "VARCHAR(255)", "—", "No", "—", "Document title (AI-generated, format: YYYY-MM-DD-FullName-DocumentType)."],
            ["description", "TEXT", "—", "Yes", "NULL", "AI-generated or user-provided document description."],
            ["file_path", "VARCHAR(255)", "—", "No", "—", "Storage path of the uploaded file."],
            ["file_hash", "VARCHAR(64)", "INDEX", "Yes", "NULL", "SHA-256 hash used for duplicate detection."],
            ["created_by", "BIGINT UNSIGNED", "FK → users(user_id), ON DELETE CASCADE", "No", "—", "User who uploaded the document."],
            ["status", "VARCHAR(255)", "INDEX", "No", "'active'", "Lifecycle state (active, archived, etc.)."],
            ["folder_id", "BIGINT UNSIGNED", "FK → folders(folder_id), ON DELETE SET NULL, INDEX", "Yes", "NULL", "Folder containing the document; NULL during AI processing (BR11)."],
            ["ai_suggested_folder", "VARCHAR(255)", "—", "Yes", "NULL", "Folder name suggested by the AI classifier."],
            ["remarks", "TEXT", "—", "Yes", "NULL", "Free-text remarks attached to the document."],
            ["physical_location", "VARCHAR(255)", "—", "Yes", "NULL", "Real-world physical filing location."],
            ["created_at", "TIMESTAMP", "INDEX", "Yes", "NULL", "Record creation timestamp."],
            ["updated_at", "TIMESTAMP", "—", "Yes", "NULL", "Record last-update timestamp."],
        ],
    },
    {
        "caption": "Table 4.4. Data Dictionary — document_embeddings",
        "intro": "Stores the 384-dimensional vector representations of document text chunks used for semantic search.",
        "rows": [
            ["embedding_id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "No", "—", "Unique identifier of the embedding record."],
            ["doc_id", "BIGINT UNSIGNED", "FK → documents(doc_id), ON DELETE CASCADE", "No", "—", "Source document of this chunk."],
            ["chunk_index", "INT", "—", "No", "—", "Sequential index of the chunk within the document."],
            ["chunk_text", "LONGTEXT", "—", "No", "—", "Raw text content of the chunk."],
            ["embedding_vector", "LONGTEXT", "—", "No", "—", "Serialized 384-dimensional embedding vector."],
            ["metadata", "JSON", "—", "Yes", "NULL", "Additional embedding metadata (page, section, etc.)."],
            ["created_at", "TIMESTAMP", "—", "No", "CURRENT_TIMESTAMP", "Record creation timestamp."],
        ],
    },
    {
        "caption": "Table 4.5. Data Dictionary — ai_conversation",
        "intro": "Stores each AI chat session initiated by a user.",
        "rows": [
            ["conversation_id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "No", "—", "Unique identifier of the conversation."],
            ["user_id", "BIGINT UNSIGNED", "FK → users(user_id), ON DELETE CASCADE, INDEX", "No", "—", "User who initiated the conversation."],
            ["doc_id", "BIGINT UNSIGNED", "INDEX", "Yes", "NULL", "Optional document anchored to the conversation."],
            ["starred", "BOOLEAN", "—", "No", "false", "Flag indicating user-pinned conversations."],
            ["started_at", "TIMESTAMP", "—", "No", "—", "Conversation start time."],
            ["ended_at", "TIMESTAMP", "—", "Yes", "NULL", "Conversation end time."],
        ],
    },
    {
        "caption": "Table 4.6. Data Dictionary — ai_histories",
        "intro": "Stores the question-and-answer entries that belong to an AI conversation.",
        "rows": [
            ["ai_history_id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "No", "—", "Unique identifier of the history entry."],
            ["conversation_id", "BIGINT UNSIGNED", "FK → ai_conversation(conversation_id), ON DELETE CASCADE", "Yes", "NULL", "Parent conversation."],
            ["doc_id", "BIGINT UNSIGNED", "FK → documents(doc_id), ON DELETE CASCADE", "Yes", "NULL", "Referenced document, if any."],
            ["user_id", "BIGINT UNSIGNED", "FK → users(user_id), ON DELETE CASCADE", "No", "—", "User who asked the question."],
            ["message_ai", "LONGTEXT", "—", "Yes", "NULL", "Raw AI message payload."],
            ["question", "LONGTEXT", "—", "No", "—", "User-submitted question."],
            ["answer", "LONGTEXT", "—", "No", "—", "AI-generated answer."],
            ["document_references", "JSON", "—", "Yes", "NULL", "Document references cited in the answer."],
            ["status", "VARCHAR(255)", "—", "No", "'completed'", "Processing status of the entry."],
            ["created_at", "TIMESTAMP", "—", "Yes", "NULL", "Entry creation timestamp."],
        ],
    },
    {
        "caption": "Table 4.7. Data Dictionary — activity_logs",
        "intro": "Stores the audit trail of user actions performed in the system.",
        "rows": [
            ["log_id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "No", "—", "Unique identifier of the log entry."],
            ["user_id", "BIGINT UNSIGNED", "FK → users(user_id), ON DELETE CASCADE", "No", "—", "User who performed the action."],
            ["doc_id", "BIGINT UNSIGNED", "FK → documents(doc_id), ON DELETE SET NULL", "Yes", "NULL", "Document affected, if applicable."],
            ["activity_type", "VARCHAR(255)", "—", "No", "—", "Type of activity (upload, edit, delete, etc.)."],
            ["activity_time", "TIMESTAMP", "—", "No", "CURRENT_TIMESTAMP", "Time the activity occurred."],
            ["activity_details", "TEXT", "—", "Yes", "NULL", "Human-readable description of the activity."],
            ["metadata", "JSON", "—", "Yes", "NULL", "Structured contextual data about the activity."],
        ],
    },
    {
        "caption": "Table 4.8. Data Dictionary — notifications",
        "intro": "Stores in-app notifications delivered to users.",
        "rows": [
            ["id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "No", "—", "Unique identifier of the notification."],
            ["user_id", "BIGINT UNSIGNED", "FK → users(user_id), ON DELETE CASCADE", "No", "—", "Recipient of the notification."],
            ["title", "VARCHAR(255)", "—", "No", "—", "Notification title."],
            ["message", "TEXT", "—", "No", "—", "Notification body."],
            ["type", "VARCHAR(255)", "—", "No", "'info'", "Notification category (info, warning, success, etc.)."],
            ["is_read", "BOOLEAN", "—", "No", "false", "Read/unread status."],
            ["created_at", "TIMESTAMP", "—", "Yes", "NULL", "Record creation timestamp."],
            ["updated_at", "TIMESTAMP", "—", "Yes", "NULL", "Record last-update timestamp."],
        ],
    },
    {
        "caption": "Table 4.9. Data Dictionary — permission_requests",
        "intro": "Stores staff-submitted requests for additional access permissions and the administrator's resolution.",
        "rows": [
            ["id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "No", "—", "Unique identifier of the request."],
            ["user_id", "BIGINT UNSIGNED", "FK → users(user_id), ON DELETE CASCADE", "No", "—", "Staff member submitting the request."],
            ["permissions", "JSON", "—", "No", "—", "Map of requested permissions (e.g., {can_view: true, can_upload: true})."],
            ["reason", "TEXT", "—", "Yes", "NULL", "Staff member's stated justification."],
            ["status", "ENUM('pending','approved','denied')", "—", "No", "'pending'", "Resolution status of the request."],
            ["admin_id", "BIGINT UNSIGNED", "FK → users(user_id), ON DELETE SET NULL", "Yes", "NULL", "Administrator who resolved the request."],
            ["admin_response", "TEXT", "—", "Yes", "NULL", "Optional response message from the administrator."],
            ["created_at", "TIMESTAMP", "—", "Yes", "NULL", "Request submission timestamp."],
            ["updated_at", "TIMESTAMP", "—", "Yes", "NULL", "Last status-change timestamp."],
        ],
    },
]


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Title
    add_heading(doc, "Data Dictionary", level=0)
    add_heading(
        doc,
        "AI-Powered Legal Document Management System",
        level=2,
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro paragraph
    add_paragraph(
        doc,
        "The data dictionary defines each table in the AI-Powered Legal Document "
        "Management System's database, including the attribute name, data type, "
        "constraints, whether the field accepts NULL, default value, and a short "
        "description of its purpose. The dictionary corresponds directly to the "
        "entities and relationships established in the Business Rules (Table 2) "
        "and the Data Model (Table 3), and reflects the database as it is "
        "implemented through Laravel migrations.",
    )

    for spec in TABLES:
        add_table_caption(doc, spec["caption"])
        add_paragraph(doc, spec["intro"], italic=True, size=10)
        add_table(doc, HEADERS, spec["rows"])

    out = r"d:\capstone\Legal_Arch_aiu\docs\Data_Dictionary.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
